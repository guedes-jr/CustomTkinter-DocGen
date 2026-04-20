import sys
import os
sys.path.append(os.path.abspath(os.curdir))

from docx import Document
from utils import docx_parser
import re

def create_split_run_template(path):
    doc = Document()
    p = doc.add_paragraph("Olá ")
    # Manually create split runs for {NOME}
    p.add_run("{")
    p.add_run("NO")
    p.add_run("ME")
    p.add_run("}")
    p.add_run("!")
    doc.save(path)

if __name__ == "__main__":
    template = "test_split_v2.docx"
    output = "test_split_v2_out.docx"
    create_split_run_template(template)
    
    doc = Document(template)
    para = doc.paragraphs[0]
    print(f"Texto original: '{para.text}'")
    print(f"Runs originais: {[r.text for r in para.runs]}")
    
    data = {"NOME": {"type": "texto", "value": "Guedes"}}
    
    # Using the FIXED parser
    docx_parser._replace_in_element(para, data)
    
    print(f"Texto após correção: '{para.text}'")
    print(f"Runs após correção: {[r.text for r in para.runs]}")
    
    if "Guedes" in para.text:
        print("SUCESSO: Variável substituída corretamente!")
    else:
        print("FALHA: Variável NÃO substituída.")
