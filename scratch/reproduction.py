from docx import Document
from docx.shared import Pt
import os

def create_sample_docx(path):
    doc = Document()
    p = doc.add_paragraph("Este é um teste com ")
    run = p.add_run("{nome}")
    run.bold = True
    run.font.size = Pt(20)
    run.font.name = 'Courier New'
    p.add_run(" e mais texto.")
    doc.save(path)

def old_replace(path, output, data):
    doc = Document(path)
    for para in doc.paragraphs:
        for key, val in data.items():
            if f"{{{key}}}" in para.text:
                para.text = para.text.replace(f"{{{key}}}", val)
    doc.save(output)

def better_replace(path, output, data):
    doc = Document(path)
    for para in doc.paragraphs:
        for run in para.runs:
            for key, val in data.items():
                if f"{{{key}}}" in run.text:
                    run.text = run.text.replace(f"{{{key}}}", val)
    doc.save(output)

if __name__ == "__main__":
    sample = "test_template.docx"
    out_old = "test_old.docx"
    out_better = "test_better.docx"
    
    create_sample_docx(sample)
    
    data = {"nome": "Antigravity"}
    
    old_replace(sample, out_old, data)
    better_replace(sample, out_better, data)
    
    print(f"Arquivos gerados: {sample}, {out_old}, {out_better}")
    # We can't easily "see" the formatting here without opening Word, 
    # but we can check if runs were preserved in out_better vs out_old.
    
    doc_old = Document(out_old)
    print(f"Old approach runs count in first para: {len(doc_old.paragraphs[0].runs)}")
    
    doc_better = Document(out_better)
    print(f"Better approach runs count in first para: {len(doc_better.paragraphs[0].runs)}")
