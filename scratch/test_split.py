from docx import Document
from docx.shared import Pt
import os
import re

def create_split_run_template(path):
    doc = Document()
    p = doc.add_paragraph("Olá ")
    # Manually create split runs for {NOME}
    p.add_run("{")
    p.add_run("NO")
    p.add_run("ME")
    p.add_run("}")
    p.add_run("!")
    doc.save(path)

def current_replace(para, data):
    pattern = r'\{([a-zA-Z0-9_]+)\}'
    if '{' not in para.text:
        return
        
    for run in para.runs:
        if '{' not in run.text:
            continue
            
        matches = re.findall(pattern, run.text)
        for key in matches:
            if key in data:
                entry = data[key]
                val = entry.get("value", "")
                run.text = run.text.replace(f"{{{key}}}", str(val))

if __name__ == "__main__":
    template = "test_split.docx"
    create_split_run_template(template)
    
    doc = Document(template)
    para = doc.paragraphs[0]
    print(f"Texto original: '{para.text}'")
    print(f"Runs: {[r.text for r in para.runs]}")
    
    data = {"NOME": {"type": "texto", "value": "Guedes"}}
    current_replace(para, data)
    
    print(f"Texto após substituição falha: '{para.text}'")
    
    # Now let's try a better logic
    def robust_replace(para, data):
        # This is a complex logic to handle split runs
        # Simplified version for demonstration: join all runs first? 
        # No, that loses formatting.
        # Let's try to find matches in para.text and then map them to runs.
        pass
