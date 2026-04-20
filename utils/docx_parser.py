from docx import Document
from docx.shared import Cm
import re
import os
import mammoth
from xhtml2pdf import pisa
import platform

def extract_variables(docx_path):
    doc = Document(docx_path)
    variables = set()
    pattern = r'\{([a-zA-Z0-9_]+)\}'
    
    for para in doc.paragraphs:
        matches = re.findall(pattern, para.text)
        for m in matches: variables.add(m)
            
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                matches = re.findall(pattern, cell.text)
                for m in matches: variables.add(m)
                    
    return list(variables)

def generate_document(template_path, output_path, data, convert_to_pdf=False):
    doc = Document(template_path)
    
    # Replace variables in paragraphs
    for para in doc.paragraphs:
        _replace_in_element(para, data)
            
    # Replace variables in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _replace_in_element(para, data)
                
    doc.save(output_path)
    
    if convert_to_pdf:
        try:
            pdf_path = output_path.replace('.docx', '.pdf')
            
            # 1. Converter Word (.docx) para HTML usando Mammoth
            with open(output_path, "rb") as docx_file:
                # O Mammoth gera um HTML limpo e incorpora imagens como base64
                result = mammoth.convert_to_html(docx_file)
                html_content = result.value
                
                # Adicionar um estilo básico para o PDF (Margens e Fontes)
                styled_html = f"""
                <html>
                <head>
                    <meta charset="UTF-8">
                    <style>
                        @page {{
                            size: a4;
                            margin: 2.5cm;
                        }}
                        body {{
                            font-family: Arial, sans-serif;
                            font-size: 11pt;
                            line-height: 1.5;
                        }}
                        table {{
                            width: 100%;
                            border-collapse: collapse;
                        }}
                        th, td {{
                            border: 1px solid #ccc;
                            padding: 8px;
                            text-align: left;
                        }}
                        img {{
                            max-width: 100%;
                            height: auto;
                        }}
                    </style>
                </head>
                <body>
                    {html_content}
                </body>
                </html>
                """
                
                # 2. Converter HTML para PDF usando xhtml2pdf (Pure Python)
                with open(pdf_path, "wb") as pdf_file:
                    pisa_status = pisa.CreatePDF(styled_html, dest=pdf_file)
                    
                if pisa_status.err:
                    raise Exception("Erro interno na geração do PDF (xhtml2pdf).")
                
            return pdf_path
        except Exception as e:
            raise Exception(f"Falha na conversão de PDF independente: {e}")
            
    return output_path

def _replace_in_element(para, data):
    pattern = r'\{([a-zA-Z0-9_]+)\}'
    
    # Identifica todas as chaves presentes no parágrafo
    matches = re.findall(pattern, para.text)
    if not matches:
        return
        
    for key in set(matches):
        if key not in data:
            continue
            
        entry = data[key]
        replacement = str(entry.get("value", ""))
        is_image = (entry.get("type") == "imagem")
        placeholder = f"{{{key}}}"
        
        # Enquanto existir o placeholder no texto do parágrafo
        while placeholder in para.text:
            # Mapeia cada caractere do parágrafo para o índice do seu Run
            char_to_run = []
            for run_idx, run in enumerate(para.runs):
                for _ in range(len(run.text)):
                    char_to_run.append(run_idx)
            
            # Encontra os índices de início e fim do placeholder no texto total
            start_char_idx = para.text.find(placeholder)
            end_char_idx = start_char_idx + len(placeholder) - 1
            
            # Descobre quais Runs compõem esse placeholder
            start_run_idx = char_to_run[start_char_idx]
            end_run_idx = char_to_run[end_char_idx]
            
            if is_image:
                # Para imagens: limpa o texto do placeholder e insere no primeiro Run
                _clear_text_range(para, start_char_idx, end_char_idx)
                if replacement and os.path.exists(replacement):
                    try:
                        para.runs[start_run_idx].add_picture(replacement, width=Cm(5))
                    except:
                        pass
            else:
                # Para texto: combina o texto dos Runs envolvidos, substitui e limpa os Runs subsequentes
                combined_text = ""
                for r_idx in range(start_run_idx, end_run_idx + 1):
                    combined_text += para.runs[r_idx].text
                
                # Substitui apenas a primeira ocorrência encontrada nesse bloco combinado
                new_combined_text = combined_text.replace(placeholder, replacement, 1)
                
                # Atualiza o primeiro Run com o texto novo
                para.runs[start_run_idx].text = new_combined_text
                
                # Limpa o texto dos outros Runs que faziam parte do placeholder
                for r_idx in range(start_run_idx + 1, end_run_idx + 1):
                    para.runs[r_idx].text = ""

def _clear_text_range(para, start, end):
    """Remove um intervalo de caracteres [start, end] através de múltiplos Runs."""
    current_idx = 0
    for run in para.runs:
        run_len = len(run.text)
        run_end = current_idx + run_len
        
        # Se este Run sobrepõe o intervalo a ser limpo
        if run_end > start and current_idx <= end:
            overlap_start = max(0, start - current_idx)
            overlap_end = min(run_len, end - current_idx + 1)
            
            text_list = list(run.text)
            del text_list[overlap_start : overlap_end]
            run.text = "".join(text_list)
            
        current_idx += run_len
