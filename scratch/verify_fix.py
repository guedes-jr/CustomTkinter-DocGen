import sys
import os
sys.path.append(os.path.abspath(os.curdir))

from utils import docx_parser
from docx import Document
from docx.shared import Pt

def create_formatted_template(path):
    doc = Document()
    p = doc.add_paragraph("Este é um teste de ")
    
    # Run 1: Bold and Large
    run = p.add_run("{negrito_grande}")
    run.bold = True
    run.font.size = Pt(24)
    run.font.name = 'Arial'
    
    p.add_run(" e agora ")
    
    # Run 2: Italic and Small
    run2 = p.add_run("{italico_pequeno}")
    run2.italic = True
    run2.font.size = Pt(8)
    run2.font.name = 'Courier New'
    
    p.add_run(" finalizado.")
    doc.save(path)

if __name__ == "__main__":
    template = "test_format_template.docx"
    output = "test_format_output.docx"
    
    create_formatted_template(template)
    
    data = {
        "negrito_grande": {"type": "texto", "value": "TEXTO GRANDE E NEGRITO"},
        "italico_pequeno": {"type": "texto", "value": "texto pequeno e itálico"}
    }
    
    docx_parser.generate_document(template, output, data)
    
    print(f"Gerado: {output}")
    
    doc = Document(output)
    para = doc.paragraphs[0]
    print(f"Número de runs final: {len(para.runs)}")
    
    for i, run in enumerate(para.runs):
        print(f"Run {i}: '{run.text}' | Bold: {run.bold} | Italic: {run.italic} | Sz: {run.font.size.pt if run.font.size else 'None'} | Name: {run.font.name}")

    # Check if formatting is preserved for the replaced parts
    # Run 1 should be the larger one
    # Run 3 should be the smaller one
    # (Runs index might be different depending on how add_run adds them, but usually 0, 1, 2, 3, 4)
    # 0: "Este é um teste de "
    # 1: "TEXTO GRANDE E NEGRITO"
    # 2: " e agora "
    # 3: "texto pequeno e itálico"
    # 4: " finalizado."
